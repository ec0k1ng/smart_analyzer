from __future__ import annotations

import html
import hashlib
import json
import math
import os
import re
import tempfile
from collections import OrderedDict
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib.parse import quote

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from jinja2 import Template
from openpyxl.utils import get_column_letter

from tcs_smart_analyzer.config.editable_configs import load_chart_view_state, load_formula_signal_definitions, load_kpi_groups
from tcs_smart_analyzer.core.features import populate_kpi_signal_values
from tcs_smart_analyzer.core.models import AnalysisResult


def _auto_fit_writer_sheets(writer) -> None:
    workbook = writer.book
    for worksheet in workbook.worksheets:
        for column_index, column_cells in enumerate(worksheet.iter_cols(), start=1):
            max_length = max((len("" if cell.value is None else str(cell.value)) for cell in column_cells), default=0)
            worksheet.column_dimensions[get_column_letter(column_index)].width = min(max(max_length + 2, 12), 60)


DEFAULT_HTML_TEMPLATE = Template(
    """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>{{ report_title }}</title>
  <style>
    :root {
      --ink: #183046;
      --muted: #5a7187;
      --line: #dbe7f1;
      --hero-a: #f8fbff;
      --hero-b: #e7f4ff;
      --group-bg: #d8eefc;
      --group-ink: #0c4a6e;
      --file-bg: #f6f9fc;
      --pass-bg: #ecfdf5;
      --pass-ink: #166534;
    --warning-bg: #fef3c7;
    --warning-ink: #92400e;
      --fail-bg: #fff1f2;
      --fail-ink: #b91c1c;
    }
    body { font-family: "Microsoft YaHei", sans-serif; margin: 0; background: linear-gradient(180deg, #f4f8fb, #eef4f7); color: var(--ink); }
    .layout { min-height: 100vh; }
    .page { width: min(1180px, calc(100vw - 404px)); margin: 0 auto; padding: 24px 0 28px; }
    .toc-sidebar { position: fixed; left: 0; top: 0; bottom: 0; width: 290px; padding: 18px 14px 18px 16px; border-right: 1px solid var(--line); border-radius: 0 26px 26px 0; background: linear-gradient(180deg, rgba(255,255,255,0.97), rgba(241,248,255,0.96)); box-shadow: 0 18px 42px rgba(24,48,70,0.14); overflow: hidden; z-index: 15; transition: transform 220ms ease, box-shadow 220ms ease; }
    .toc-toggle { position: fixed; left: -9999px; }
    .toc-fab { position: fixed; left: 306px; top: 16px; display: inline-flex; align-items: center; gap: 8px; padding: 10px 14px; border-radius: 999px; border: 1px solid #c9dceb; background: rgba(255,255,255,0.96); color: #0f4c81; box-shadow: 0 12px 28px rgba(24,48,70,0.12); cursor: pointer; z-index: 20; font-size: 13px; font-weight: 800; transition: left 220ms ease, transform 220ms ease; }
    .toc-fab:hover { transform: translateY(-1px); }
    .toc-fab-close { display: inline; }
    .toc-fab-open { display: none; }
    .toc-header { display: flex; align-items: center; justify-content: space-between; gap: 10px; padding: 16px 18px; background: linear-gradient(180deg, #f8fbff, #edf6ff); color: #0f4c81; font-weight: 800; font-size: 16px; border-radius: 18px; }
    .toc-badge { font-size: 12px; color: var(--muted); font-weight: 700; }
    .toc-nav { padding: 14px 8px 18px 4px; max-height: calc(100vh - 118px); overflow: auto; }
    .toc-group { margin-bottom: 14px; }
    .toc-group-link { display: block; text-decoration: none; color: #0f4c81; font-weight: 800; padding: 8px 10px; border-radius: 12px; background: #edf6ff; border: 1px solid #cfe2ef; }
    .toc-file-list { list-style: none; margin: 8px 0 0 0; padding: 0 0 0 10px; }
    .toc-file-list li { margin: 6px 0; }
    .toc-file-link { display: block; text-decoration: none; color: #31485d; padding: 6px 10px; border-radius: 10px; }
    .toc-file-link:hover, .toc-group-link:hover { background: #dff0ff; }
    .toc-toggle:not(:checked) ~ .toc-sidebar { transform: translateX(-100%); box-shadow: none; }
    .toc-toggle:not(:checked) ~ .toc-fab { left: 14px; }
    .toc-toggle:not(:checked) ~ .toc-fab .toc-fab-close { display: none; }
    .toc-toggle:not(:checked) ~ .toc-fab .toc-fab-open { display: inline; }
    .hero { padding: 22px 24px; border: 1px solid #cfe2ef; border-radius: 22px; background: linear-gradient(135deg, var(--hero-a), var(--hero-b)); box-shadow: 0 12px 32px rgba(24,48,70,0.08); margin-bottom: 22px; }
    .hero h1 { margin: 0 0 12px 0; font-size: 30px; color: #0c4a6e; }
    .hero-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(210px, 1fr)); gap: 12px; }
    .hero-card { background: rgba(255,255,255,0.7); border: 1px solid rgba(207,226,239,0.9); border-radius: 16px; padding: 12px 14px; }
    .hero-card .label { font-size: 12px; color: var(--muted); margin-bottom: 4px; }
    .hero-card .value { font-weight: 700; color: var(--ink); }
    .section-title { margin: 26px 0 12px 0; font-size: 18px; color: #0f4c81; }
    .group-section { margin-bottom: 28px; scroll-margin-top: 24px; }
    .group-heading { margin: 0 0 12px 0; padding: 12px 16px; border-radius: 16px; background: var(--group-bg); color: var(--group-ink); font-size: 18px; font-weight: 800; }
    .result-card { margin-bottom: 20px; border: 1px solid var(--line); border-radius: 18px; overflow: hidden; background: #fff; box-shadow: 0 8px 20px rgba(15, 76, 129, 0.05); }
    .file-bar { display: flex; justify-content: space-between; gap: 16px; padding: 12px 16px; background: var(--file-bg); color: #334155; border-top: 1px solid var(--line); }
    .file-name { font-weight: 700; }
    .file-summary { color: var(--muted); font-size: 13px; }
    .chart-preview { padding: 14px 16px 8px 16px; background: #ffffff; }
    .chart-preview img { width: 100%; border: 1px solid #dbe7f1; border-radius: 14px; background: linear-gradient(180deg, #fcfdff, #f5f9fc); }
    table { border-collapse: collapse; width: 100%; table-layout: fixed; }
    th, td { border-top: 1px solid var(--line); padding: 9px 10px; font-size: 13px; vertical-align: top; }
    th { background: #f7fbff; text-align: left; color: #23415c; }
    td.col-unit, th.col-unit { width: 88px; text-align: center; }
    td.col-value, th.col-value { width: 120px; text-align: right; }
    td.col-result, th.col-result { width: 104px; text-align: center; }
    .result-pill { display: inline-block; min-width: 72px; padding: 4px 10px; border-radius: 999px; font-weight: 800; }
    .result-pill.pass { background: var(--pass-bg); color: var(--pass-ink); }
    .result-pill.warning { background: var(--warning-bg); color: var(--warning-ink); }
    .result-pill.fail { background: var(--fail-bg); color: var(--fail-ink); }
        @media (max-width: 1080px) {
            .page { width: auto; margin: 70px 18px 0 18px; padding: 0 0 24px; }
            .toc-sidebar { width: min(310px, calc(100vw - 34px)); top: 58px; bottom: 18px; border-radius: 22px; left: 12px; }
            .toc-fab { left: 14px; top: 12px; }
            .toc-nav { max-height: calc(100vh - 170px); }
        }
  </style>
</head>
<body>
    <div class="layout">
        {% if grouped_results %}
        <input id="toc-toggle" class="toc-toggle" type="checkbox" checked>
        <label class="toc-fab" for="toc-toggle"><span class="toc-fab-close">隐藏目录</span><span class="toc-fab-open">显示目录</span></label>
        <aside class="toc-sidebar">
            <div class="toc-header"><span>目录</span><span class="toc-badge">KPI 分组 / 数据名称</span></div>
            <nav class="toc-nav">
                {% for group in grouped_results %}
                <div class="toc-group">
                    <a class="toc-group-link" href="#{{ group['group_anchor'] }}">{{ group['group_name'] }}</a>
                    <ul class="toc-file-list">
                        {% for item in group['items'] %}
                        <li><a class="toc-file-link" href="#{{ item['file_anchor'] }}">{{ item.source_name }}</a></li>
                        {% endfor %}
                    </ul>
                </div>
                {% endfor %}
            </nav>
        </aside>
        {% endif %}
    <div class="page">
    <section class="hero">
      <h1>{{ report_title }}</h1>
      <div class="hero-grid">
        <div class="hero-card"><div class="label">汇总文件数</div><div class="value">{{ file_count }}</div></div>
        <div class="hero-card"><div class="label">分析配置</div><div class="value">{{ analysis_profile }}</div></div>
        <div class="hero-card"><div class="label">分析摘要</div><div class="value">{{ report_summary }}</div></div>
        <div class="hero-card"><div class="label">生成时间</div><div class="value">{{ generated_at }}</div></div>
      </div>
    </section>
    <div class="section-title">分析结果</div>
        {% for group in grouped_results %}
        <section class="group-section" id="{{ group['group_anchor'] }}">
            <div class="group-heading">{{ group['group_name'] }}</div>
            {% for item in group['items'] %}
            <section class="result-card" id="{{ item['file_anchor'] }}">
                <div class="file-bar">
                    <div class="file-name">{{ item.source_name }}</div>
                    <div class="file-summary">KPI 项数：{{ item.kpi_count }}</div>
                </div>
                {% if item.chart_preview_data_uri %}
                <div class="chart-preview">
                    <img src="{{ item.chart_preview_data_uri }}" alt="{{ item.source_name }} 曲线工作表1预览">
                </div>
                {% endif %}
                <table>
                    <thead>
                        <tr>
                            <th>KPI</th>
                            <th>描述</th>
                            <th class="col-unit">单位</th>
                            <th class="col-value">数值</th>
                            <th>达标要求</th>
                            <th class="col-result">结果</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for kpi in item.kpis %}
                        <tr>
                            <td>{{ kpi.title }}</td>
                            <td>{{ kpi.description }}</td>
                            <td class="col-unit">{{ kpi.unit }}</td>
                            <td class="col-value">{{ kpi.value_text }}</td>
                            <td>{{ kpi.rule_description }}</td>
                            <td class="col-result"><span class="result-pill {{ kpi.status }}">{{ kpi.result_label }}</span></td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </section>
            {% endfor %}
        </section>
    {% endfor %}
  </div>
    </div>
</body>
</html>
"""
)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_rfonts(target, font_name: str) -> None:  # noqa: ANN001
    r_fonts = target.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        target.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), font_name)


def _apply_run_style(run, *, font_name: str = "Microsoft YaHei", size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:  # noqa: ANN001
    run.font.name = font_name
    _set_rfonts(run._element.get_or_add_rPr(), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _apply_style_font(style, *, font_name: str = "Microsoft YaHei", size: float = 10.5) -> None:  # noqa: ANN001
    style.font.name = font_name
    style.font.size = Pt(size)
    _set_rfonts(style._element.get_or_add_rPr(), font_name)


def _style_cell_paragraphs(cell, *, size: float = 10.5, bold: bool | None = None) -> None:  # noqa: ANN001
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            _apply_run_style(run, size=size, bold=bold)


def _append_bookmark(paragraph, bookmark_name: str, bookmark_id: int) -> None:  # noqa: ANN001
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bookmark_id))
    bookmark_start.set(qn("w:name"), bookmark_name)
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bookmark_id))
    insert_index = 0
    for index, child in enumerate(paragraph._p):
        if child.tag == qn("w:pPr"):
            insert_index = index + 1
            break
    paragraph._p.insert(insert_index, bookmark_start)
    paragraph._p.append(bookmark_end)


def _append_internal_hyperlink(paragraph, target_anchor: str, text: str, *, color: RGBColor, bold: bool = False, left_indent: float | None = None) -> None:  # noqa: ANN001
    if left_indent is not None:
        paragraph.paragraph_format.left_indent = Pt(left_indent)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), target_anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    _set_rfonts(run_pr, "Microsoft YaHei")
    color_element = OxmlElement("w:color")
    color_hex = "".join(f"{int(component):02X}" for component in color)
    color_element.set(qn("w:val"), color_hex)
    run_pr.append(color_element)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    run_pr.append(underline)
    if bold:
        run_pr.append(OxmlElement("w:b"))
    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), str(int(round((11.5 if bold else 10.5) * 2))))
    run_pr.append(size_element)
    run.append(run_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _insert_word_toc(document: Document, grouped_results: list[dict[str, Any]]) -> None:
    toc_heading = document.add_paragraph(style="Heading 1")
    toc_run = toc_heading.add_run("目录")
    _apply_run_style(toc_run, size=15, bold=True, color=RGBColor(12, 74, 110))
    for group in grouped_results:
        group_paragraph = document.add_paragraph()
        group_paragraph.paragraph_format.space_after = Pt(3)
        _append_internal_hyperlink(
            group_paragraph,
            str(group.get("group_anchor", "")),
            str(group.get("group_name", "")),
            color=RGBColor(12, 74, 110),
            bold=True,
        )
        for item in group.get("items", []):
            file_paragraph = document.add_paragraph()
            file_paragraph.paragraph_format.space_after = Pt(1)
            _append_internal_hyperlink(
                file_paragraph,
                str(item.get("file_anchor", "")),
                f"- {item.get('source_name', '')}",
                color=RGBColor(51, 65, 85),
                left_indent=16,
            )


def _group_name_lookup() -> dict[str, str]:
    groups = {"__all_kpis__": "默认组（全部 KPI）"}
    groups.update({str(item.get("key", "__all_kpis__")): str(item.get("name", "默认组（全部 KPI）")) for item in load_kpi_groups()})
    return groups


def _safe_anchor_id(prefix: str, *parts: str) -> str:
    raw = "|".join(str(part) for part in parts if str(part))
    summary = re.sub(r"[^a-zA-Z0-9]+", "_", raw).strip("_").lower()[:8] or prefix
    digest = hashlib.sha1(raw.encode("utf-8")).hexdigest()[:12]
    return f"{prefix}_{summary}_{digest}"


def _group_anchor_id(group_key: str, group_name: str) -> str:
    return _safe_anchor_id("grp", group_key, group_name)


def _file_anchor_id(group_key: str, source_name: str, source_path: str) -> str:
    return _safe_anchor_id("file", group_key, source_name, source_path)


def _prepare_chart_preview_frame(result: AnalysisResult, related_results: list[AnalysisResult] | None = None) -> pd.DataFrame:
    frame_data = result.normalized_frame.copy()
    for related_result in related_results or [result]:
        populate_kpi_signal_values(
            related_result.normalized_frame,
            related_result.kpis,
            group_key=str(related_result.context.metadata.get("kpi_group_key", "__all_kpis__")),
        )
        for item in related_result.kpis:
            if item.signal_values is not None:
                frame_data[item.name] = item.signal_values.reindex(frame_data.index)
            else:
                frame_data[item.name] = float(item.value)
    local_context: dict[str, object] = {column: frame_data[column] for column in frame_data.columns}
    for definition in load_formula_signal_definitions():
        expression = str(definition.get("expression", "")).strip()
        name = str(definition.get("name", "")).strip()
        if not expression or not name:
            continue
        try:
            value = eval(expression, {"__builtins__": {}}, {**local_context, "math": math, "min": min, "max": max, "abs": abs})
        except Exception:
            continue
        if hasattr(value, "reindex"):
            frame_data[name] = value.reindex(frame_data.index)
        else:
            frame_data[name] = value
        local_context[name] = frame_data[name]
    return frame_data


def _chart_preview_panel_signals(frame: pd.DataFrame) -> list[list[str]]:
    available = {str(column).strip() for column in frame.columns if str(column).strip() and str(column).strip() != "time_s"}
    view_state = load_chart_view_state()
    sheets = view_state.get("sheets", []) if isinstance(view_state, dict) else []
    if sheets and isinstance(sheets[0], dict):
        panels: list[list[str]] = []
        for panel in sheets[0].get("panels", []):
            if not isinstance(panel, dict):
                continue
            hidden = {str(name).strip() for name in panel.get("hidden_signals", []) if str(name).strip()}
            signals = [str(name).strip() for name in panel.get("signals", []) if str(name).strip() in available and str(name).strip() not in hidden]
            if signals:
                panels.append(signals)
        if panels:
            return panels

    priority = [
        "vehicle_speed_kph",
        "wheel_speed_fl_kph",
        "wheel_speed_fr_kph",
        "wheel_speed_rl_kph",
        "wheel_speed_rr_kph",
        "slip_kph",
        "tcs_active",
    ]
    fallback = [[signal_name] for signal_name in priority if signal_name in available]
    if fallback:
        return fallback
    return [[signal_name] for signal_name in list(available)]


def _chart_preview_tick_values(min_value: float, max_value: float, steps: int = 4) -> list[float]:
    if steps <= 0:
        return [min_value, max_value]
    span = max_value - min_value
    if span <= 1e-9:
        return [min_value for _ in range(steps + 1)]
    return [min_value + span * index / steps for index in range(steps + 1)]


def _chart_preview_tick_label(value: float) -> str:
    magnitude = abs(value)
    if magnitude >= 1000:
        return f"{value:.0f}"
    if magnitude >= 100:
        return f"{value:.1f}"
    if magnitude >= 1:
        return f"{value:.2f}"
    return f"{value:.3f}"


def _sample_signal_frame(frame: pd.DataFrame, signal_name: str, max_points: int = 240) -> pd.DataFrame:
    signal_frame = frame[["time_s", signal_name]].copy()
    signal_frame["time_s"] = pd.to_numeric(signal_frame["time_s"], errors="coerce")
    signal_frame[signal_name] = pd.to_numeric(signal_frame[signal_name], errors="coerce")
    signal_frame = signal_frame.dropna()
    if signal_frame.empty or len(signal_frame) <= max_points:
        return signal_frame
    step = max(1, len(signal_frame) // max_points)
    return signal_frame.iloc[::step].copy()


def _chart_preview_svg(result: AnalysisResult, related_results: list[AnalysisResult] | None = None) -> str:
    frame = _prepare_chart_preview_frame(result, related_results=related_results)
    if frame.empty or "time_s" not in frame.columns:
        return ""

    time_series = pd.to_numeric(frame["time_s"], errors="coerce").dropna()
    if time_series.empty:
        return ""

    panel_signals = _chart_preview_panel_signals(frame)
    if not panel_signals:
        return ""

    width = 1040
    outer_left = 22
    left = 94
    right = 28
    panel_height = 188
    gap = 18
    top_padding = 18
    bottom_padding = 34
    plot_width = width - left - right
    height = top_padding + len(panel_signals) * panel_height + max(0, len(panel_signals) - 1) * gap + bottom_padding
    x_min = float(time_series.min())
    x_max = float(time_series.max())
    x_span = max(x_max - x_min, 1e-9)
    palette = ["#0f766e", "#2563eb", "#c2410c", "#b91c1c", "#7c3aed", "#0891b2"]

    svg_parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<rect width="100%" height="100%" fill="#f8fbff" rx="18" ry="18"/>',
    ]

    for panel_index, signals in enumerate(panel_signals):
        panel_top = top_padding + panel_index * (panel_height + gap)
        plot_top = panel_top + 42
        plot_height = panel_height - 72
        svg_parts.append(f'<rect x="12" y="{panel_top}" width="{width - 24}" height="{panel_height}" rx="14" ry="14" fill="#ffffff" stroke="#dbe7f1"/>')

        numeric_values: list[float] = []
        sampled_frames: list[tuple[str, pd.DataFrame]] = []
        for signal_name in signals:
            if signal_name not in frame.columns:
                continue
            sampled = _sample_signal_frame(frame, signal_name)
            if sampled.empty:
                continue
            sampled_frames.append((signal_name, sampled))
            numeric_values.extend(sampled[signal_name].astype(float).tolist())

        if not numeric_values:
            continue

        y_min = min(numeric_values)
        y_max = max(numeric_values)
        y_span = y_max - y_min
        if y_span <= 1e-9:
            padding = max(abs(y_min) * 0.15, 1.0)
            y_min -= padding
            y_max += padding
        else:
            padding = max(y_span * 0.08, 1e-6)
            y_min -= padding
            y_max += padding
        y_span = max(y_max - y_min, 1e-9)

        legend_x = 24
        legend_y = panel_top + 18
        for color_index, signal_name in enumerate(signals):
            color = palette[color_index % len(palette)]
            block_x = legend_x + color_index * 150
            svg_parts.append(f'<line x1="{block_x}" y1="{legend_y - 4}" x2="{block_x + 18}" y2="{legend_y - 4}" stroke="{color}" stroke-width="3" stroke-linecap="round"/>')
            svg_parts.append(f'<text x="{block_x + 24}" y="{legend_y}" font-size="12" font-weight="700" fill="{color}">{html.escape(signal_name)}</text>')

        y_ticks = _chart_preview_tick_values(y_min, y_max)
        for tick_value in y_ticks:
            tick_ratio = (tick_value - y_min) / y_span
            y = plot_top + plot_height - tick_ratio * plot_height
            svg_parts.append(f'<line x1="{left}" y1="{y:.2f}" x2="{left + plot_width}" y2="{y:.2f}" stroke="#edf2f7" stroke-width="1"/>')
            svg_parts.append(f'<line x1="{left - 5}" y1="{y:.2f}" x2="{left}" y2="{y:.2f}" stroke="#64748b" stroke-width="1"/>')
            svg_parts.append(f'<text x="{left - 10}" y="{y + 4:.2f}" text-anchor="end" font-size="10.5" fill="#475569">{html.escape(_chart_preview_tick_label(tick_value))}</text>')

        x_ticks = _chart_preview_tick_values(x_min, x_max)
        for tick_value in x_ticks:
            tick_ratio = (tick_value - x_min) / x_span
            x = left + tick_ratio * plot_width
            svg_parts.append(f'<line x1="{x:.2f}" y1="{plot_top}" x2="{x:.2f}" y2="{plot_top + plot_height}" stroke="#f1f5f9" stroke-width="1"/>')
            svg_parts.append(f'<line x1="{x:.2f}" y1="{plot_top + plot_height}" x2="{x:.2f}" y2="{plot_top + plot_height + 5}" stroke="#64748b" stroke-width="1"/>')
            svg_parts.append(f'<text x="{x:.2f}" y="{plot_top + plot_height + 18}" text-anchor="middle" font-size="10.5" fill="#475569">{html.escape(_chart_preview_tick_label(tick_value))}</text>')

        svg_parts.append(f'<line x1="{left}" y1="{plot_top}" x2="{left}" y2="{plot_top + plot_height}" stroke="#64748b" stroke-width="1.3"/>')
        svg_parts.append(f'<line x1="{left}" y1="{plot_top + plot_height}" x2="{left + plot_width}" y2="{plot_top + plot_height}" stroke="#64748b" stroke-width="1.3"/>')
        svg_parts.append(f'<text x="{outer_left}" y="{plot_top + plot_height / 2:.2f}" transform="rotate(-90 {outer_left},{plot_top + plot_height / 2:.2f})" text-anchor="middle" font-size="10.5" font-weight="700" fill="#475569">数值</text>')
        svg_parts.append(f'<text x="{left + plot_width / 2:.2f}" y="{plot_top + plot_height + 30}" text-anchor="middle" font-size="10.5" font-weight="700" fill="#475569">时间 (s)</text>')

        for color_index, (signal_name, sampled) in enumerate(sampled_frames):
            points: list[str] = []
            for time_value, signal_value in zip(sampled["time_s"], sampled[signal_name]):
                x = left + ((float(time_value) - x_min) / x_span) * plot_width
                y = plot_top + plot_height - ((float(signal_value) - y_min) / y_span) * plot_height
                points.append(f"{x:.2f},{y:.2f}")
            if points:
                color = palette[color_index % len(palette)]
                svg_parts.append(f'<polyline fill="none" stroke="{color}" stroke-width="2" points="{" ".join(points)}"/>')

    svg_parts.append("</svg>")
    return "".join(svg_parts)


def _chart_preview_data_uri(result: AnalysisResult, related_results: list[AnalysisResult] | None = None) -> str:
    svg = _chart_preview_svg(result, related_results=related_results)
    if not svg:
        return ""
    return "data:image/svg+xml;charset=utf-8," + quote(svg)


_QT_GUI_APP = None


def _ensure_qt_gui_application() -> bool:
    global _QT_GUI_APP

    try:
        from PySide6.QtGui import QGuiApplication
    except ImportError:
        return False

    if QGuiApplication.instance() is not None:
        return True

    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    _QT_GUI_APP = QGuiApplication(["smart-analyzer-report"])
    return True


def _chart_preview_png_path(result: AnalysisResult, related_results: list[AnalysisResult] | None = None) -> Path | None:
    svg = _chart_preview_svg(result, related_results=related_results)
    if not svg:
        return None

    if not _ensure_qt_gui_application():
        return None

    try:
        from PySide6.QtCore import QByteArray
        from PySide6.QtGui import QImage, QPainter
        from PySide6.QtSvg import QSvgRenderer
    except ImportError:
        return None

    renderer = QSvgRenderer(QByteArray(svg.encode("utf-8")))
    if not renderer.isValid():
        return None

    default_size = renderer.defaultSize()
    width = max(int(default_size.width() or 1040), 1040)
    height = max(int(default_size.height() or 520), 520)
    image = QImage(width, height, QImage.Format.Format_ARGB32)
    image.fill(0xFFFFFFFF)

    painter = QPainter(image)
    renderer.render(painter)
    painter.end()

    handle = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    handle.close()
    png_path = Path(handle.name)
    if not image.save(str(png_path), "PNG"):
        png_path.unlink(missing_ok=True)
        return None
    return png_path


def _format_value(value: float | None) -> str:
    if value is None:
        return "-"
    return f"{float(value):.4f}"


def _serialize_kpi(item) -> dict[str, Any]:  # noqa: ANN001
    return {
        "name": item.name,
        "title": item.title,
        "value": item.value,
        "value_text": _format_value(item.value),
        "unit": item.unit,
        "description": item.description,
        "rule_description": item.rule_description,
        "status": item.status,
        "result_label": item.result_label,
    }


def _serialize_rule(item) -> dict[str, Any]:  # noqa: ANN001
    return {
        "rule_id": item.rule_id,
        "title": item.title,
        "status": item.status,
        "measured_value": item.measured_value,
        "threshold_value": item.threshold_value,
        "message": item.message,
    }


def _serialize_result(
    result: AnalysisResult,
    report_title: str,
    group_names: dict[str, str],
    related_results_by_path: dict[str, list[AnalysisResult]] | None = None,
) -> dict[str, Any]:
    group_key = str(result.context.metadata.get("kpi_group_key", "__all_kpis__"))
    related_results = None if related_results_by_path is None else related_results_by_path.get(str(result.context.source_path), [result])
    serialized = {
        "report_title": report_title,
        "source_path": str(result.context.source_path),
        "source_name": result.context.source_path.name,
        "source_stem": result.context.source_path.stem,
        "analysis_profile": result.context.metadata.get("analysis_profile", "default"),
        "generated_at": result.context.metadata.get("generated_at", ""),
        "mapped_columns": result.context.mapped_columns,
        "metadata": result.context.metadata,
        "group_key": group_key,
        "group_name": group_names.get(group_key, group_key),
        "group_anchor": _group_anchor_id(group_key, group_names.get(group_key, group_key)),
        "file_anchor": _file_anchor_id(group_key, result.context.source_path.name, str(result.context.source_path)),
        "kpi_count": len(result.kpis),
        "chart_preview_data_uri": _chart_preview_data_uri(result, related_results=related_results),
        "kpis": [_serialize_kpi(item) for item in result.kpis],
        "rules": [_serialize_rule(item) for item in result.rule_results],
    }
    return serialized


def build_report_context(result: AnalysisResult, report_title: str = "TCS 打滑控制自动分析报告") -> dict[str, Any]:
    group_names = _group_name_lookup()
    related_results_by_path = {str(result.context.source_path): [result]}
    serialized = _serialize_result(result, report_title, group_names, related_results_by_path=related_results_by_path)
    grouped_results = [{
        "group_key": serialized["group_key"],
        "group_name": serialized["group_name"],
        "group_anchor": serialized["group_anchor"],
        "items": [serialized],
    }]
    context = dict(serialized)
    context["report_title"] = report_title
    context["file_count"] = 1
    context["report_summary"] = f"共输出 {len(serialized['kpis'])} 条 KPI 结果。"
    context["kpis"] = list(serialized["kpis"])
    context["rules"] = list(serialized["rules"])
    context["metadata"] = dict(serialized["metadata"])
    context["mapped_columns"] = dict(serialized["mapped_columns"])
    context["results"] = [serialized]
    context["files"] = context["results"]
    context["grouped_results"] = grouped_results
    return context


def _flatten_batch_records(results: list[AnalysisResult]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    kpis: list[dict[str, Any]] = []
    rules: list[dict[str, Any]] = []
    group_names = _group_name_lookup()
    related_results_by_path = _related_results_by_path(results)
    for result in results:
        serialized = _serialize_result(result, "", group_names, related_results_by_path=related_results_by_path)
        for item in serialized["kpis"]:
            kpis.append({
                "source_path": serialized["source_path"],
                "source_name": serialized["source_name"],
                "source_stem": serialized["source_stem"],
                "group_key": serialized["group_key"],
                "group_name": serialized["group_name"],
                **item,
            })
        for item in serialized["rules"]:
            rules.append({
                "source_path": serialized["source_path"],
                "source_name": serialized["source_name"],
                "group_key": serialized["group_key"],
                "group_name": serialized["group_name"],
                **item,
            })
    return kpis, rules


def _related_results_by_path(results: list[AnalysisResult]) -> dict[str, list[AnalysisResult]]:
    lookup: dict[str, list[AnalysisResult]] = {}
    for result in results:
        lookup.setdefault(str(result.context.source_path), []).append(result)
    return lookup


def build_batch_report_context(results: list[AnalysisResult], report_title: str = "TCS 打滑控制自动分析报告") -> dict[str, Any]:
    if not results:
        generated_at = datetime.now().astimezone().isoformat(timespec="seconds")
        return {
            "report_title": report_title,
            "generated_at": generated_at,
            "analysis_profile": "default",
            "source_path": "",
            "source_name": "",
            "source_stem": "",
            "report_summary": "",
            "mapped_columns": {},
            "metadata": {},
            "kpis": [],
            "rules": [],
            "results": [],
            "files": [],
            "grouped_results": [],
            "file_count": 0,
        }

    group_names = _group_name_lookup()
    related_results_by_path = _related_results_by_path(results)
    items = [
        _serialize_result(result, report_title=report_title, group_names=group_names, related_results_by_path=related_results_by_path)
        for result in results
    ]
    kpis, rules = _flatten_batch_records(results)
    grouped_lookup: OrderedDict[str, dict[str, Any]] = OrderedDict()
    for item in items:
        group_key = str(item.get("group_key", "__all_kpis__"))
        if group_key not in grouped_lookup:
            grouped_lookup[group_key] = {
                "group_key": group_key,
                "group_name": item.get("group_name", group_key),
                "group_anchor": item.get("group_anchor", _group_anchor_id(group_key, str(item.get("group_name", group_key)))),
                "items": [],
            }
        grouped_lookup[group_key]["items"].append(item)
    profiles = sorted({str(item.get("analysis_profile", "default")) for item in items})
    generated_at = max((str(item.get("generated_at", "")) for item in items), default="") or datetime.now().astimezone().isoformat(timespec="seconds")
    report_summary = f"共分析 {len(results)} 个文件，输出 {len(kpis)} 条 KPI 结果。"
    return {
        "report_title": report_title,
        "generated_at": generated_at,
        "analysis_profile": ", ".join(profiles),
        "source_path": "多文件汇总报告",
        "source_name": "",
        "source_stem": "",
        "report_summary": report_summary,
        "mapped_columns": {},
        "metadata": {
            "analysis_profiles": profiles,
            "file_count": len(results),
            "config_paths": sorted({str(item["metadata"].get("config_path")) for item in items if item["metadata"].get("config_path")}),
            "group_keys": sorted({str(item["metadata"].get("kpi_group_key", "__all_kpis__")) for item in items}),
        },
        "kpis": kpis,
        "rules": rules,
        "results": items,
        "files": items,
        "grouped_results": list(grouped_lookup.values()),
        "file_count": len(results),
    }


def batch_report_filename(report_title: str) -> str:
    stem = re.sub(r"[^a-zA-Z0-9_\-\u4e00-\u9fff]+", "_", report_title.strip()).strip("_")
    return f"{stem or 'tcs_batch_report'}_summary.html"


def batch_word_filename(report_title: str) -> str:
    stem = re.sub(r"[^a-zA-Z0-9_\-\u4e00-\u9fff]+", "_", report_title.strip()).strip("_")
    return f"{stem or 'tcs_batch_report'}_summary.docx"


def export_excel(result: AnalysisResult, output_path: str | Path) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    kpi_frame = pd.DataFrame(
        [
            {
                "name": item.name,
                "title": item.title,
                "value": item.value,
                "unit": item.unit,
                "description": item.description,
                "rule_description": item.rule_description,
                "status": item.status,
                "assessment_message": item.assessment_message,
                "threshold_value": item.threshold_value,
                "threshold_source": item.threshold_source,
                "pass_condition": item.pass_condition,
                "result_label": item.result_label,
            }
            for item in result.kpis
        ]
    )
    rule_frame = pd.DataFrame(
        [
            {
                "rule_id": item.rule_id,
                "category": item.category,
                "title": item.title,
                "status": item.status,
                "severity": item.severity,
                "measured_value": item.measured_value,
                "threshold_value": item.threshold_value,
                "unit": item.unit,
                "message": item.message,
                "threshold_source": item.threshold_source,
                "confidence": item.confidence,
            }
            for item in result.rule_results
        ]
    )

    with pd.ExcelWriter(output) as writer:
        kpi_frame.to_excel(writer, sheet_name="kpis", index=False)
        rule_frame.to_excel(writer, sheet_name="rules", index=False)
        result.normalized_frame.to_excel(writer, sheet_name="normalized_data", index=False)
        _auto_fit_writer_sheets(writer)

    return output


def export_html(
    result: AnalysisResult | list[AnalysisResult],
    output_path: str | Path,
    template_path: str | Path | None = None,
    report_title: str = "TCS 打滑控制自动分析报告",
) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    template = DEFAULT_HTML_TEMPLATE
    if template_path is not None:
        template = Template(Path(template_path).read_text(encoding="utf-8"))

    if isinstance(result, list):
        context = build_batch_report_context(result, report_title=report_title)
    else:
        context = build_report_context(result, report_title=report_title)
    html = template.render(**context)
    output.write_text(html, encoding="utf-8")
    return output


def export_word(
    result: AnalysisResult | list[AnalysisResult],
    output_path: str | Path,
    report_title: str = "TCS 打滑控制自动分析报告",
) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    context = build_batch_report_context(result, report_title=report_title) if isinstance(result, list) else build_report_context(result, report_title=report_title)
    result_lookup = {
        str(item.context.source_path): item
        for item in (result if isinstance(result, list) else [result])
    }
    related_results_by_path = _related_results_by_path(result if isinstance(result, list) else [result])
    document = Document()
    normal_style = document.styles["Normal"]
    _apply_style_font(normal_style)
    for style_name, size in [("Heading 1", 15), ("Heading 2", 12)]:
        if style_name in document.styles:
            _apply_style_font(document.styles[style_name], size=size)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(report_title)
    _apply_run_style(run, size=20, bold=True, color=RGBColor(12, 74, 110))

    meta_table = document.add_table(rows=2, cols=2)
    meta_table.style = "Table Grid"
    meta_pairs = [
        ("汇总文件数", str(context.get("file_count", 0))),
        ("分析配置", str(context.get("analysis_profile", "default"))),
        ("分析摘要", str(context.get("report_summary", ""))),
        ("生成时间", str(context.get("generated_at", ""))),
    ]
    for index, (label, value) in enumerate(meta_pairs):
        row = index // 2
        col = (index % 2) * 1
        cell = meta_table.cell(row, col)
        cell.text = f"{label}: {value}"
        _set_cell_shading(cell, "E7F4FF")
        _style_cell_paragraphs(cell, size=10.5, bold=False)

    document.add_paragraph("")
    _insert_word_toc(document, context.get("grouped_results", []))
    document.add_page_break()
    bookmark_id = 1
    for group in context.get("grouped_results", []):
        group_paragraph = document.add_paragraph(style="Heading 1")
        _append_bookmark(group_paragraph, str(group.get("group_anchor", "group")), bookmark_id)
        bookmark_id += 1
        group_run = group_paragraph.add_run(str(group.get("group_name", "")))
        _apply_run_style(group_run, size=13, bold=True, color=RGBColor(12, 74, 110))

        for result_item in group.get("items", []):
            file_paragraph = document.add_paragraph(style="Heading 2")
            _append_bookmark(file_paragraph, str(result_item.get("file_anchor", "result-file")), bookmark_id)
            bookmark_id += 1
            file_run = file_paragraph.add_run(str(result_item.get("source_name", "")))
            _apply_run_style(file_run, size=11.5, bold=True)
            count_run = file_paragraph.add_run(f"    KPI 项数: {result_item.get('kpi_count', 0)}")
            _apply_run_style(count_run, size=11.5, bold=True)

            preview_path = None
            source_path = str(result_item.get("source_path", ""))
            matched_result = result_lookup.get(source_path)
            if matched_result is not None:
                preview_path = _chart_preview_png_path(matched_result, related_results=related_results_by_path.get(source_path, [matched_result]))

            if preview_path is not None:
                preview_paragraph = document.add_paragraph()
                preview_run = preview_paragraph.add_run()
                preview_run.add_picture(str(preview_path), width=Pt(460))
                preview_path.unlink(missing_ok=True)

            table = document.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            table.autofit = True
            headers = ["KPI", "描述", "单位", "数值", "达标要求", "结果"]
            for index, header in enumerate(headers):
                cell = table.rows[0].cells[index]
                cell.text = header
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                _set_cell_shading(cell, "D8EEFC")
                _style_cell_paragraphs(cell, size=10.5, bold=True)
            for item in result_item.get("kpis", []):
                row = table.add_row().cells
                row[0].text = str(item.get("title", ""))
                row[1].text = str(item.get("description", ""))
                row[2].text = str(item.get("unit", ""))
                row[3].text = str(item.get("value_text", "-"))
                row[4].text = str(item.get("rule_description", ""))
                row[5].text = str(item.get("result_label", ""))
                status = str(item.get("status", "pass"))
                if status == "pass":
                    _set_cell_shading(row[5], "ECFDF5")
                elif status == "warning":
                    _set_cell_shading(row[5], "FEF3C7")
                else:
                    _set_cell_shading(row[5], "FFF1F2")
                for cell in row:
                    _style_cell_paragraphs(cell, size=10.5, bold=False)
            document.add_paragraph("")

    document.save(output)
    return output


def export_json_summary(result: AnalysisResult, output_path: str | Path) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    time_start = None
    time_end = None
    if not result.normalized_frame.empty and "time_s" in result.normalized_frame.columns:
        time_start = float(result.normalized_frame["time_s"].min())
        time_end = float(result.normalized_frame["time_s"].max())

    payload = {
        "source_path": str(result.context.source_path),
        "generated_at": result.context.metadata.get("generated_at"),
        "analyzer_version": result.context.metadata.get("analyzer_version"),
        "analysis_profile": result.context.metadata.get("analysis_profile", "default"),
        "config_path": result.context.metadata.get("config_path"),
        "settings_metadata": result.context.metadata.get("settings_metadata", {}),
        "mapped_columns": result.context.mapped_columns,
        "rule_settings": result.context.metadata.get("rule_settings", {}),
        "data_overview": {
            "row_count": int(len(result.normalized_frame)),
            "time_start_s": time_start,
            "time_end_s": time_end,
        },
        "kpis": [
            {
                "name": item.name,
                "value": item.value,
                "unit": item.unit,
                "description": item.description,
            }
            for item in result.kpis
        ],
        "rules": [
            {
                "rule_id": item.rule_id,
                "category": item.category,
                "title": item.title,
                "status": item.status,
                "severity": item.severity,
                "measured_value": item.measured_value,
                "threshold_value": item.threshold_value,
                "unit": item.unit,
                "message": item.message,
                "threshold_source": item.threshold_source,
                "confidence": item.confidence,
            }
            for item in result.rule_results
        ],
    }

    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return output
