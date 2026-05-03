from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path
from urllib.parse import unquote

from docx import Document

from tcs_smart_analyzer.config.editable_configs import get_config_file_paths, save_interface_signal_tables
from tcs_smart_analyzer.core.engine import AnalysisEngine
from tcs_smart_analyzer.reporting.exporters import _file_anchor_id, _group_anchor_id, build_batch_report_context, build_report_context, export_html, export_json_summary, export_word


class ExporterTests(unittest.TestCase):
    def setUp(self) -> None:
        config_paths = get_config_file_paths()
        self._interface_mapping_path = config_paths["interface_mapping"]
        self._interface_mapping_backup = self._interface_mapping_path.read_bytes() if self._interface_mapping_path.exists() else None
        save_interface_signal_tables(
            [
                {"standard_signal": "time_s", "actual_names": ["time_s"]},
                {"standard_signal": "wheel_speed_fl_kph", "actual_names": ["wheel_speed_fl"]},
                {"standard_signal": "wheel_speed_fr_kph", "actual_names": ["wheel_speed_fr"]},
                {"standard_signal": "wheel_speed_rl_kph", "actual_names": ["wheel_speed_rl"]},
                {"standard_signal": "wheel_speed_rr_kph", "actual_names": ["wheel_speed_rr"]},
                {"standard_signal": "vehicle_speed_kph", "actual_names": ["vehicle_speed"]},
                {"standard_signal": "accel_pedal_pct", "actual_names": ["accel_pedal_pct"]},
                {"standard_signal": "brake_depth_pct", "actual_names": ["time_s*0"]},
                {"standard_signal": "torque_request_nm", "actual_names": ["torque_request_nm"]},
                {"standard_signal": "torque_actual_nm", "actual_names": ["torque_actual_nm"]},
                {"standard_signal": "longitudinal_accel_mps2", "actual_names": ["longitudinal_accel_mps2"]},
                {"standard_signal": "abs_active_fl", "actual_names": ["time_s*0"]},
                {"standard_signal": "abs_active_fr", "actual_names": ["time_s*0"]},
                {"standard_signal": "abs_active_rl", "actual_names": ["time_s*0"]},
                {"standard_signal": "abs_active_rr", "actual_names": ["time_s*0"]},
                {"standard_signal": "yaw_rate_degps", "actual_names": ["time_s*0"]},
                {"standard_signal": "steering_wheel_angle_deg", "actual_names": ["time_s*0"]},
                {"standard_signal": "tcs_active", "actual_names": ["tcs_active"]},
                {"standard_signal": "tcs_active_fl", "actual_names": ["tcs_active"]},
                {"standard_signal": "tcs_active_fr", "actual_names": ["tcs_active"]},
                {"standard_signal": "tcs_active_rl", "actual_names": ["tcs_active"]},
                {"standard_signal": "tcs_active_rr", "actual_names": ["tcs_active"]},
            ],
            [],
        )

    def tearDown(self) -> None:
        if self._interface_mapping_backup is None:
            if self._interface_mapping_path.exists():
                self._interface_mapping_path.unlink()
            return
        self._interface_mapping_path.write_bytes(self._interface_mapping_backup)

    def test_json_summary_contains_traceability_metadata(self) -> None:
        project_root = Path(__file__).resolve().parents[1]
        demo_file = project_root / "sample_data" / "tcs_demo.csv"

        engine = AnalysisEngine()
        result = engine.analyze_file(demo_file)

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "summary.json"
            export_json_summary(result, output_path)
            payload = json.loads(output_path.read_text(encoding="utf-8"))

        self.assertEqual(payload["source_path"], str(demo_file))
        self.assertEqual(payload["analysis_profile"], "default")
        self.assertIn("rule_settings", payload)
        self.assertGreaterEqual(len(payload["rules"]), 4)
        self.assertGreater(payload["data_overview"]["row_count"], 0)
        self.assertNotIn("scenarios", payload)

    def test_html_export_accepts_custom_template(self) -> None:
        project_root = Path(__file__).resolve().parents[1]
        demo_file = project_root / "sample_data" / "tcs_demo.csv"

        engine = AnalysisEngine()
        result = engine.analyze_file(demo_file)

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "custom_template.html"
            output_path = temp_path / "report.html"
            template_path.write_text("<html><body><h1>{{ report_title }}</h1><p>{{ source_path }}</p></body></html>", encoding="utf-8")

            export_html(result, output_path, template_path=template_path, report_title="自定义模板报告")
            content = output_path.read_text(encoding="utf-8")

        self.assertIn("自定义模板报告", content)
        self.assertIn(str(demo_file), content)

    def test_example_template_renders_chart_preview(self) -> None:
        project_root = Path(__file__).resolve().parents[1]
        demo_file = project_root / "sample_data" / "tcs_demo.csv"

        engine = AnalysisEngine()
        result = engine.analyze_file(demo_file)

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "report.html"
            template_path = project_root / "src" / "tcs_smart_analyzer" / "config" / "report_templates" / "00_example_and_guide.html"
            export_html(result, output_path, template_path=template_path, report_title="示例模板报告")
            content = output_path.read_text(encoding="utf-8")

        self.assertIn("data:image/svg+xml", content)
        self.assertIn("曲线工作表1预览", content)
        self.assertIn("toc-sidebar", content)
        self.assertIn("toc-fab", content)
        self.assertIn("隐藏目录", content)
        self.assertIn("显示目录", content)
        self.assertIn("KPI 分组", content)
        self.assertIn("数据名称", content)

    def test_report_context_exposes_file_name_variables(self) -> None:
        project_root = Path(__file__).resolve().parents[1]
        demo_file = project_root / "sample_data" / "tcs_demo.csv"

        engine = AnalysisEngine()
        result = engine.analyze_file(demo_file)
        context = build_report_context(result)

        self.assertEqual(context["source_path"], str(demo_file))
        self.assertEqual(context["source_name"], demo_file.name)
        self.assertEqual(context["source_stem"], demo_file.stem)
        self.assertEqual(context["file_count"], 1)
        self.assertEqual(context["results"][0]["group_name"], "默认组（全部 KPI）")
        self.assertIn("report_summary", context)
        self.assertIn("value_text", context["results"][0]["kpis"][0])
        self.assertTrue(context["results"][0]["chart_preview_data_uri"].startswith("data:image/svg+xml"))

        svg = unquote(context["results"][0]["chart_preview_data_uri"].split(",", 1)[1])
        self.assertIn("时间 (s)", svg)
        self.assertIn("数值", svg)
        self.assertGreaterEqual(svg.count('stroke="#dbe7f1"'), 5)

    def test_batch_report_context_flattens_multiple_results(self) -> None:
        project_root = Path(__file__).resolve().parents[1]
        demo_file = project_root / "sample_data" / "tcs_demo.csv"

        engine = AnalysisEngine()
        first = engine.analyze_file(demo_file)
        second = engine.analyze_file(demo_file)
        context = build_batch_report_context([first, second])

        self.assertEqual(context["file_count"], 2)
        self.assertEqual(len(context["results"]), 2)
        self.assertGreaterEqual(len(context["kpis"]), len(first.kpis) * 2)
        self.assertIn("group_name", context["kpis"][0])

    def test_word_export_matches_report_structure(self) -> None:
        project_root = Path(__file__).resolve().parents[1]
        demo_file = project_root / "sample_data" / "tcs_demo.csv"

        engine = AnalysisEngine()
        result = engine.analyze_file(demo_file)

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "report.docx"
            export_word(result, output_path, report_title="Word 报告")
            document = Document(output_path)

        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Word 报告", text)
        self.assertIn(demo_file.name, text)
        self.assertGreaterEqual(len(document.tables), 2)
        self.assertIn("目录", text)
        self.assertIn("默认组（全部 KPI）", text)
        heading_styles = {paragraph.style.name for paragraph in document.paragraphs if paragraph.text}
        self.assertIn("Heading 1", heading_styles)
        self.assertIn("Heading 2", heading_styles)
        self.assertNotIn("右键更新域以生成目录", text)
        self.assertIn("w:hyperlink", document._element.xml)
        self.assertIn("w:bookmarkStart", document._element.xml)
        self.assertIn('w:anchor="file_', document._element.xml)
        self.assertIn('w:anchor="grp_', document._element.xml)
        self.assertNotIn('w:u w:val="single"', document._element.xml)

    def test_word_export_keeps_file_heading_run_style_consistent(self) -> None:
        project_root = Path(__file__).resolve().parents[1]
        demo_file = project_root / "sample_data" / "tcs_demo.csv"

        engine = AnalysisEngine()
        result = engine.analyze_file(demo_file)

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "report.docx"
            export_word(result, output_path, report_title="Word 报告")
            document = Document(output_path)

        target_paragraph = next(paragraph for paragraph in document.paragraphs if "KPI 项数" in paragraph.text)
        self.assertGreaterEqual(len(target_paragraph.runs), 2)
        font_names = {run.font.name for run in target_paragraph.runs}
        font_sizes = {None if run.font.size is None else round(run.font.size.pt, 1) for run in target_paragraph.runs}
        bold_values = {run.bold for run in target_paragraph.runs}

        self.assertEqual(font_names, {"Microsoft YaHei"})
        self.assertEqual(font_sizes, {11.5})
        self.assertEqual(bold_values, {True})

    def test_word_export_places_bookmark_after_paragraph_properties(self) -> None:
        project_root = Path(__file__).resolve().parents[1]
        demo_file = project_root / "sample_data" / "tcs_demo.csv"

        engine = AnalysisEngine()
        result = engine.analyze_file(demo_file)

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "report.docx"
            export_word(result, output_path, report_title="Word 报告")
            document = Document(output_path)

        target_paragraph = next(paragraph for paragraph in document.paragraphs if "KPI 项数" in paragraph.text)
        paragraph_xml = target_paragraph._p.xml
        self.assertIn("w:pPr", paragraph_xml)
        self.assertIn("w:bookmarkStart", paragraph_xml)
        self.assertLess(paragraph_xml.index("w:pPr"), paragraph_xml.index("w:bookmarkStart"))

    def test_word_anchor_ids_stay_within_safe_length(self) -> None:
        group_anchor = _group_anchor_id("__all_kpis__", "默认组（全部 KPI）")
        file_anchor = _file_anchor_id("__all_kpis__", "tcs_demo.csv", "sample_data/tcs_demo.csv")

        self.assertLessEqual(len(group_anchor), 40)
        self.assertLessEqual(len(file_anchor), 40)
        self.assertTrue(group_anchor.startswith("grp_"))
        self.assertTrue(file_anchor.startswith("file_"))


if __name__ == "__main__":
    unittest.main()